from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
)
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# COLORS
# ============================================================

ILAW_GREEN = "00B050"
LIGHT_GREEN = "D9EAD3"
MEDIUM_GREEN = "93D977"
WHITE = "FFFFFF"
BLACK = "000000"


# ============================================================
# MAIN DOCX GENERATOR
# ============================================================

def generate_ilaw_docx(lesson_plan: dict) -> BytesIO:

    document = Document()

    section = document.sections[0]

    section.orientation = WD_ORIENT.LANDSCAPE

    section.page_width = Inches(11)
    section.page_height = Inches(8.5)

    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    # --------------------------------------------------------
    # DEFAULT FONT
    # --------------------------------------------------------

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)

    # --------------------------------------------------------
    # DATA
    # --------------------------------------------------------

    info = lesson_plan.get(
        "lesson_information",
        {},
    )

    references = lesson_plan.get(
        "references",
        [],
    )

    declaration = lesson_plan.get(
        "declaration_of_ai_use",
        "",
    )

    intentions = lesson_plan.get(
        "intentions",
        {},
    )

    experiences = lesson_plan.get(
        "learning_experiences",
        {},
    )

    sessions = lesson_plan.get(
        "sessions",
        [],
    )

    assessment = lesson_plan.get(
        "assessment",
        {},
    )

    ways_forward = lesson_plan.get(
        "ways_forward",
        {},
    )

    prepared = lesson_plan.get(
        "prepared_checked_noted",
        {},
    )

    # ========================================================
    # TITLE
    # ========================================================

    title_table = document.add_table(
        rows=1,
        cols=1,
    )

    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_table.autofit = False

    title_cell = title_table.cell(0, 0)

    shade_cell(
        title_cell,
        ILAW_GREEN,
    )

    set_cell_text(
        title_cell,
        "Lesson Plan - ILAW Format",
        bold=True,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    set_cell_width(
        title_cell,
        7.57,
    )

    # Small spacing
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    # ========================================================
    # LESSON INFORMATION
    # ========================================================

    add_section_header(
        document,
        "LESSON INFORMATION",
    )

    info_table = document.add_table(
        rows=6,
        cols=2,
    )

    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False

    info_rows = [
        (
            "Lesson Title",
            info.get("title", ""),
        ),
        (
            "Learning Area/s",
            info.get("learning_area", ""),
        ),
        (
            "Name of Teacher/s",
            ", ".join(
                info.get("teachers", [])
            ),
        ),
        (
            "Grade Level and Section",
            build_grade_section(info),
        ),
        (
            "No. of Sessions",
            str(
                info.get(
                    "sessions",
                    len(sessions),
                )
            ),
        ),
        (
            "Sessions",
            "",
        ),
    ]

    for row_index, (label, value) in enumerate(
        info_rows
    ):

        left = info_table.cell(
            row_index,
            0,
        )

        right = info_table.cell(
            row_index,
            1,
        )

        shade_cell(
            left,
            LIGHT_GREEN,
        )

        set_cell_text(
            left,
            label,
            bold=True,
            size=8,
        )

        if label == "Sessions":

            add_session_headers(
                right,
                sessions,
            )

        else:

            set_cell_text(
                right,
                value,
                size=9,
            )

        set_cell_width(
            left,
            1.85,
        )

        set_cell_width(
            right,
            5.72,
        )

    set_table_borders(
        info_table,
    )

    # ========================================================
    # REFERENCES
    # ========================================================

    add_section_header(
        document,
        "References (books, websites, toolkits, etc.)",
    )

    references_table = document.add_table(
        rows=1,
        cols=2,
    )

    references_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    references_table.autofit = False

    ref_left = references_table.cell(0, 0)
    ref_right = references_table.cell(0, 1)

    shade_cell(
        ref_left,
        LIGHT_GREEN,
    )

    set_cell_text(
        ref_left,
        "References (books,\nwebsites, toolkits,\netc.)",
        bold=False,
        size=8,
    )

    set_cell_width(
        ref_left,
        1.85,
    )

    set_cell_width(
        ref_right,
        5.72,
    )

    add_list_to_cell(
        ref_right,
        references,
    )

    set_table_borders(
        references_table,
    )

    # ========================================================
    # DECLARATION OF AI USE
    # ========================================================

    ai_table = document.add_table(
        rows=1,
        cols=2,
    )

    ai_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    ai_table.autofit = False

    ai_left = ai_table.cell(0, 0)
    ai_right = ai_table.cell(0, 1)

    shade_cell(
        ai_left,
        LIGHT_GREEN,
    )

    set_cell_text(
        ai_left,
        "Declaration of AI Use",
        bold=True,
        size=8,
    )

    set_cell_width(
        ai_left,
        1.85,
    )

    set_cell_width(
        ai_right,
        5.72,
    )

    set_cell_text(
        ai_right,
        declaration,
        size=8,
    )

    set_table_borders(
        ai_table,
    )

    # ========================================================
    # I - INTENTIONS
    # ========================================================

    add_section_header(
        document,
        "I - INTENTIONS",
        subtitle=(
            "Meaningful learning experiences are anchored "
            "in how we frame them. Start by deciding what "
            "you want learners to master by the end of "
            "the lesson."
        ),
    )

    intentions_table = document.add_table(
        rows=1,
        cols=2,
    )

    intentions_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    intentions_table.autofit = False

    left = intentions_table.cell(0, 0)
    right = intentions_table.cell(0, 1)

    shade_cell(
        left,
        LIGHT_GREEN,
    )

    set_cell_text(
        left,
        "Learning Competencies\n\n"
        "Write the competencies from the curriculum "
        "that are targeted, and the content or "
        "performance standards applicable to the sessions.",
        size=7,
    )

    set_cell_width(
        left,
        1.85,
    )

    set_cell_width(
        right,
        5.72,
    )

    add_label_value(
        right,
        "Content Standard",
        intentions.get(
            "content_standard",
            "",
        ),
    )

    add_label_value(
        right,
        "Performance Standard",
        intentions.get(
            "performance_standard",
            "",
        ),
    )

    add_list_label_value(
        right,
        "Learning Competencies",
        intentions.get(
            "learning_competencies",
            [],
        ),
    )

    add_list_label_value(
        right,
        "Specific Objectives",
        intentions.get(
            "specific_objectives",
            [],
        ),
    )

    add_label_value(
        right,
        "Learning Objectives",
        intentions.get(
            "learning_objectives",
            "",
        ),
    )

    add_label_value(
        right,
        "Learner Context",
        intentions.get(
            "learner_context",
            "",
        ),
    )

    set_table_borders(
        intentions_table,
    )

    # ========================================================
    # L - LEARNING EXPERIENCES
    # ========================================================

    add_section_header(
        document,
        "L - LEARNING EXPERIENCES",
        subtitle=(
            "A learning experience is like a thoughtfully "
            "designed journey. Each activity and interaction "
            "builds towards meaningful understanding and growth."
        ),
    )

    # Learning Resources
    learning_resource_table = document.add_table(
        rows=1,
        cols=2,
    )

    learning_resource_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    learning_resource_table.autofit = False

    left = learning_resource_table.cell(0, 0)
    right = learning_resource_table.cell(0, 1)

    shade_cell(
        left,
        LIGHT_GREEN,
    )

    set_cell_text(
        left,
        "Learning Resources",
        bold=True,
        size=8,
    )

    set_cell_text(
        right,
        experiences.get(
            "learning_resources",
            "",
        ),
        size=9,
    )

    set_cell_width(
        left,
        1.85,
    )

    set_cell_width(
        right,
        5.72,
    )

    set_table_borders(
        learning_resource_table,
    )

    # Pre-Lesson
    pre_table = document.add_table(
        rows=1,
        cols=2,
    )

    pre_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    pre_table.autofit = False

    left = pre_table.cell(0, 0)
    right = pre_table.cell(0, 1)

    shade_cell(
        left,
        LIGHT_GREEN,
    )

    set_cell_text(
        left,
        "Pre-Lesson",
        bold=True,
        size=8,
    )

    set_cell_text(
        right,
        experiences.get(
            "pre_lesson",
            "",
        ),
        size=8,
    )

    set_cell_width(
        left,
        1.85,
    )

    set_cell_width(
        right,
        5.72,
    )

    set_table_borders(
        pre_table,
    )

    # Flow / Daylong
    flow = experiences.get(
        "flow_daylong",
        {},
    )

    flow_table = document.add_table(
        rows=1,
        cols=2,
    )

    flow_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    flow_table.autofit = False

    left = flow_table.cell(0, 0)
    right = flow_table.cell(0, 1)

    shade_cell(
        left,
        LIGHT_GREEN,
    )

    set_cell_text(
        left,
        "Flow / Daylong",
        bold=True,
        size=8,
    )

    add_label_value(
        right,
        "Activity",
        flow.get(
            "activity",
            "",
        ),
    )

    add_label_value(
        right,
        "Discussion",
        flow.get(
            "discussion",
            "",
        ),
    )

    add_label_value(
        right,
        "Deduction / Generalization",
        flow.get(
            "deduction",
            "",
        ),
    )

    add_list_label_value(
        right,
        "Concepts",
        flow.get(
            "concepts",
            [],
        ),
    )

    set_cell_width(
        left,
        1.85,
    )

    set_cell_width(
        right,
        5.72,
    )

    set_table_borders(
        flow_table,
    )

    # Integration
    integration_table = document.add_table(
        rows=1,
        cols=2,
    )

    integration_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    integration_table.autofit = False

    left = integration_table.cell(0, 0)
    right = integration_table.cell(0, 1)

    shade_cell(
        left,
        LIGHT_GREEN,
    )

    set_cell_text(
        left,
        "Opportunities for Integration",
        bold=True,
        size=8,
    )

    set_cell_text(
        right,
        experiences.get(
            "opportunities_for_integration",
            "",
        ),
        size=8,
    )

    set_cell_width(
        left,
        1.85,
    )

    set_cell_width(
        right,
        5.72,
    )

    set_table_borders(
        integration_table,
    )
    # ========================================================
    # LEARNING SESSIONS
    # ========================================================

    add_section_header(
        document,
        "LEARNING SESSIONS",
    )

    if not sessions:
        sessions = []

    for index, session in enumerate(sessions):

        session_number = session.get(
            "session_number",
            index + 1,
        )

        # ----------------------------------------------------
        # Keep each session together when possible
        # ----------------------------------------------------

        session_table = document.add_table(
            rows=4,
            cols=2,
        )

        session_table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )

        session_table.autofit = False

        # ----------------------------------------------------
        # SESSION HEADER
        # ----------------------------------------------------

        header_left = session_table.cell(
            0,
            0,
        )

        header_right = session_table.cell(
            0,
            1,
        )

        shade_cell(
            header_left,
            MEDIUM_GREEN,
        )

        shade_cell(
            header_right,
            MEDIUM_GREEN,
        )

        set_cell_text(
            header_left,
            f"Session {session_number}",
            bold=True,
            size=8,
        )

        set_cell_text(
            header_right,
            session.get(
                "topic",
                "",
            ),
            bold=True,
            size=8,
        )

        # ----------------------------------------------------
        # LEARNING ACTIVITIES
        # ----------------------------------------------------

        activity_label = session_table.cell(
            1,
            0,
        )

        activity_value = session_table.cell(
            1,
            1,
        )

        shade_cell(
            activity_label,
            LIGHT_GREEN,
        )

        set_cell_text(
            activity_label,
            "Learning Activities",
            bold=True,
            size=7,
        )

        set_cell_text(
            activity_value,
            session.get(
                "activities",
                "",
            ),
            size=8,
        )

        # ----------------------------------------------------
        # ASSESSMENT
        # ----------------------------------------------------

        assessment_label = session_table.cell(
            2,
            0,
        )

        assessment_value = session_table.cell(
            2,
            1,
        )

        shade_cell(
            assessment_label,
            LIGHT_GREEN,
        )

        set_cell_text(
            assessment_label,
            "Assessment",
            bold=True,
            size=7,
        )

        set_cell_text(
            assessment_value,
            session.get(
                "assessment",
                "",
            ),
            size=8,
        )

        # ----------------------------------------------------
        # ADDITIONAL DETAILS
        # ----------------------------------------------------

        details_label = session_table.cell(
            3,
            0,
        )

        details_value = session_table.cell(
            3,
            1,
        )

        shade_cell(
            details_label,
            LIGHT_GREEN,
        )

        set_cell_text(
            details_label,
            "Additional Details",
            bold=True,
            size=7,
        )

        set_cell_text(
            details_value,
            session.get(
                "details",
                "",
            ),
            size=8,
        )

        # ----------------------------------------------------
        # WIDTHS
        # ----------------------------------------------------

        for row in session_table.rows:

            set_cell_width(
                row.cells[0],
                1.85,
            )

            set_cell_width(
                row.cells[1],
                5.72,
            )

        set_table_borders(
            session_table,
        )

        # ----------------------------------------------------
        # SPACE BETWEEN SESSIONS
        # ----------------------------------------------------

        if index < len(sessions) - 1:

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
    # ========================================================
    # A - ASSESSMENT
    # ========================================================

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)

    add_section_header(
        document,
        "A - ASSESSMENT",
        subtitle=(
            "Assessments reveal what learners have gained "
            "and what they still need help with."
        ),
    )

    assessment_table = document.add_table(
        rows=1,
        cols=2,
    )

    assessment_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    assessment_table.autofit = False

    left = assessment_table.cell(0, 0)
    right = assessment_table.cell(0, 1)

    shade_cell(
        left,
        LIGHT_GREEN,
    )

    set_cell_text(
        left,
        "Formative Assessment",
        bold=True,
        size=8,
    )

    add_label_value(
        right,
        "Formative Assessment",
        assessment.get(
            "formative_assessment",
            "",
        ),
    )

    add_list_label_value(
        right,
        "Guide Questions",
        assessment.get(
            "guide_questions",
            [],
        ),
    )

    for row in assessment_table.rows:

        set_cell_width(
            row.cells[0],
            1.85,
        )

        set_cell_width(
            row.cells[1],
            5.72,
        )

    set_table_borders(
        assessment_table,
    )

    # ========================================================
    # W - WAYS FORWARD
    # ========================================================

    add_section_header(
        document,
        "W - WAYS FORWARD",
        subtitle=(
            "Meaningful learning can happen beyond the "
            "classroom - for both the learners and teacher."
        ),
    )

    ways_table = document.add_table(
        rows=3,
        cols=2,
    )

    ways_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    ways_table.autofit = False

    ways_rows = [
        (
            "Extended Learning",
            ways_forward.get(
                "extended_learning",
                "",
            ),
        ),
        (
            "Reflections",
            ways_forward.get(
                "reflections",
                "",
            ),
        ),
        (
            "Application",
            ways_forward.get(
                "application",
                "",
            ),
        ),
    ]

    for index, (label, value) in enumerate(
        ways_rows
    ):

        left = ways_table.cell(
            index,
            0,
        )

        right = ways_table.cell(
            index,
            1,
        )

        shade_cell(
            left,
            LIGHT_GREEN,
        )

        set_cell_text(
            left,
            label,
            bold=True,
            size=8,
        )

        set_cell_text(
            right,
            value,
            size=8,
        )

        set_cell_width(
            left,
            1.85,
        )

        set_cell_width(
            right,
            5.72,
        )

    set_table_borders(
        ways_table,
    )

    # ========================================================
    # PREPARED / CHECKED / NOTED
    # ========================================================

    add_section_header(
        document,
        "PREPARED, CHECKED AND NOTED",
    )

    sign_table = document.add_table(
        rows=2,
        cols=3,
    )

    sign_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    sign_table.autofit = False

    labels = [
        "Prepared by:",
        "Checked by:",
        "Noted:",
    ]

    values = [
        prepared.get(
            "prepared_by",
            "",
        ),
        prepared.get(
            "checked_by",
            "",
        ),
        prepared.get(
            "noted_by",
            "",
        ),
    ]

    for column in range(3):

        header = sign_table.cell(
            0,
            column,
        )

        body = sign_table.cell(
            1,
            column,
        )

        shade_cell(
            header,
            WHITE,
        )

        set_cell_text(
            header,
            labels[column],
            bold=True,
            size=8,
        )

        set_cell_text(
            body,
            values[column],
            bold=True,
            size=8,
        )

        set_cell_width(
            header,
            2.52,
        )

        set_cell_width(
            body,
            2.52,
        )

    set_table_borders(
        sign_table,
    )

    # ========================================================
    # FINALIZE
    # ========================================================

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output


# ============================================================
# HELPERS
# ============================================================

def build_grade_section(info):

    grade = info.get(
        "grade_level",
        "",
    )

    section = info.get(
        "section",
        "",
    )

    if grade and section:
        return f"{grade} - {section}"

    return grade or section


def add_section_header(
    document,
    title,
    subtitle=None,
    width=7.57,
):

    table = document.add_table(
        rows=1,
        cols=1,
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.autofit = False

    cell = table.cell(0, 0)

    shade_cell(
        cell,
        ILAW_GREEN,
    )

    cell.text = ""

    paragraph = cell.paragraphs[0]

    run = paragraph.add_run(
        title
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    if subtitle:

        run = paragraph.add_run(
            "\n" + subtitle
        )

        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(7)

    set_cell_width(
        cell,
        width,
    )

    set_table_width(
        table,
        width,
    )

    set_table_borders(
        table,
    )


def add_small_spacing(document):

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def add_session_headers(cell, sessions):

    cell.text = ""

    if not sessions:
        return

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)

    for index, session in enumerate(sessions):

        if index > 0:
            paragraph.add_run("     ")

        run = paragraph.add_run(
            f"Session {session.get('session_number', index + 1)}"
        )

        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(7)

    shade_cell(
        cell,
        MEDIUM_GREEN,
    )


def add_label_value(
    cell,
    label,
    value,
):

    paragraph = cell.add_paragraph()

    paragraph.paragraph_format.space_after = Pt(2)

    run = paragraph.add_run(
        f"{label}: "
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8)

    run = paragraph.add_run(
        str(value or "")
    )

    run.font.name = "Arial"
    run.font.size = Pt(8)


def add_list_label_value(
    cell,
    label,
    items,
):

    paragraph = cell.add_paragraph()

    paragraph.paragraph_format.space_after = Pt(1)

    run = paragraph.add_run(
        f"{label}:"
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8)

    for item in items or []:

        bullet = cell.add_paragraph(
            style="List Bullet"
        )

        bullet.paragraph_format.left_indent = (
            Inches(0.15)
        )

        bullet.paragraph_format.space_after = Pt(0)

        run = bullet.add_run(
            str(item)
        )

        run.font.name = "Arial"
        run.font.size = Pt(8)


def add_list_to_cell(
    cell,
    items,
):

    cell.text = ""

    if not items:

        return

    for item in items:

        paragraph = cell.add_paragraph(
            style="List Bullet"
        )

        paragraph.paragraph_format.space_after = Pt(0)

        run = paragraph.add_run(
            str(item)
        )

        run.font.name = "Arial"
        run.font.size = Pt(8)


def set_cell_text(
    cell,
    text,
    bold=False,
    size=9,
    align=None,
):

    cell.text = ""

    paragraph = cell.paragraphs[0]

    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)

    if align is not None:
        paragraph.alignment = align

    run = paragraph.add_run(
        str(text or "")
    )

    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


def set_cell_width(
    cell,
    width,
):

    cell.width = Inches(width)

    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(
        qn("w:tcW")
    )

    if tc_w is None:

        tc_w = OxmlElement(
            "w:tcW"
        )

        tc_pr.append(tc_w)

    tc_w.set(
        qn("w:w"),
        str(
            int(
                width * 1440
            )
        ),
    )

    tc_w.set(
        qn("w:type"),
        "dxa",
    )


def shade_cell(
    cell,
    color,
):

    tc_pr = cell._tc.get_or_add_tcPr()

    shd = tc_pr.find(
        qn("w:shd")
    )

    if shd is None:

        shd = OxmlElement(
            "w:shd"
        )

        tc_pr.append(shd)

    shd.set(
        qn("w:fill"),
        color,
    )


def set_table_width(
    table,
    width=7.57,
):

    """Force the entire table to a fixed width in inches."""

    table.autofit = False

    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(
        qn("w:tblW")
    )

    if tbl_w is None:

        tbl_w = OxmlElement(
            "w:tblW"
        )

        tbl_pr.append(tbl_w)

    tbl_w.set(
        qn("w:w"),
        str(
            int(
                width * 1440
            )
        ),
    )

    tbl_w.set(
        qn("w:type"),
        "dxa",
    )


def set_table_borders(
    table,
    color=BLACK,
    size="6",
):

    # Keep every ILAW table aligned to the same 7.57-inch content width.
    set_table_width(
        table,
        7.57,
    )

    tbl = table._tbl

    tbl_pr = tbl.tblPr

    borders = tbl_pr.first_child_found_in(
        "w:tblBorders"
    )

    if borders is None:

        borders = OxmlElement(
            "w:tblBorders"
        )

        tbl_pr.append(borders)

    for edge in (
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ):

        tag = f"w:{edge}"

        element = borders.find(
            qn(tag)
        )

        if element is None:

            element = OxmlElement(
                tag
            )

            borders.append(element)

        element.set(
            qn("w:val"),
            "single",
        )

        element.set(
            qn("w:sz"),
            size,
        )

        element.set(
            qn("w:space"),
            "0",
        )

        element.set(
            qn("w:color"),
            color,
        )


# ============================================================
# TEST / SAFETY
# ============================================================

if __name__ == "__main__":

    print(
        "ILAW DOCX generator loaded successfully."
    )