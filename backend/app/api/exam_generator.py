import json
import re

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from io import BytesIO

from docx import Document
from docx.shared import Pt

from app.services.llm import generate_response


router = APIRouter()


class ExamRequest(BaseModel):
    prompt: str


class ExamResponse(BaseModel):
    title: str
    instructions: str
    questions: list
    answer_key: list


@router.post("/generate", response_model=ExamResponse)
async def generate_exam(request: ExamRequest):

    if not request.prompt.strip():
        raise HTTPException(
            status_code=400,
            detail="Please enter an exam prompt."
        )

    prompt = f"""
You are Mak-AI, an educational exam generator.

Create an exam based on this teacher request:

{request.prompt}

Requirements:

1. Follow the requested grade level and subject.
2. Follow the requested number of items.
3. If the teacher does not specify a number of items,
   create 20 items.
4. Make the questions appropriate for the requested grade.
5. Avoid duplicate questions.
6. Make questions clear and classroom-ready.
7. Use multiple-choice questions unless the teacher
   specifically requests another format.
8. Each multiple-choice question must have exactly
   four choices: A, B, C, D.
9. Include the correct answer.
10. Include a short explanation for every answer.
11. Create a clear exam title.
12. Create simple exam instructions.

Return ONLY valid JSON.

Use this exact structure:

{{
  "title": "Exam Title",
  "instructions": "Exam instructions",
  "questions": [
    {{
      "number": 1,
      "question": "Question text",
      "choices": {{
        "A": "Choice A",
        "B": "Choice B",
        "C": "Choice C",
        "D": "Choice D"
      }},
      "answer": "A",
      "explanation": "Short explanation"
    }}
  ]
}}

Do not return markdown.
Do not return ```json.
Do not add text outside the JSON.
"""

    try:

        result = generate_response([
            {
                "role": "system",
                "content": (
                    "You are a reliable educational exam generator. "
                    "Return strict JSON when requested."
                ),
            },
            {
                "role": "user",
                "content": prompt,
            },
        ])

        if not isinstance(result, str):
            result = str(result)

        result = result.strip()

        result = re.sub(
            r"^```json\s*",
            "",
            result,
            flags=re.IGNORECASE,
        )

        result = re.sub(
            r"^```\s*",
            "",
            result,
        )

        result = re.sub(
            r"\s*```$",
            "",
            result,
        )

        data = json.loads(result)

        questions = data.get("questions", [])

        if not questions:
            raise ValueError("No questions were generated.")

        return {
            "title": data.get(
                "title",
                "Mak-AI Generated Exam"
            ),
            "instructions": data.get(
                "instructions",
                "Choose the best answer."
            ),
            "questions": questions,
            "answer_key": [
                {
                    "number": q.get("number", index + 1),
                    "answer": q.get("answer", ""),
                    "explanation": q.get(
                        "explanation",
                        ""
                    ),
                }
                for index, q in enumerate(questions)
            ],
        }

    except json.JSONDecodeError as error:

        print("Exam JSON error:", error)

        raise HTTPException(
            status_code=500,
            detail="Mak-AI returned an invalid exam format."
        )

    except Exception as error:

        print("Exam generation error:", error)

        raise HTTPException(
            status_code=500,
            detail="Unable to generate the exam."
        )


@router.post("/download")
async def download_exam(exam: ExamResponse):

    document = Document()

    # Title
    title = document.add_paragraph()

    title_run = title.add_run(exam.title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    title.alignment = 1

    document.add_paragraph()

    # Student information
    document.add_paragraph(
        "Name: ______________________________________________"
    )

    document.add_paragraph(
        "Section: _____________________________________________"
    )

    document.add_paragraph(
        "Date: _________________________________________________"
    )

    document.add_paragraph()

    # Instructions
    instruction_title = document.add_paragraph()

    instruction_run = instruction_title.add_run(
        "Instructions"
    )

    instruction_run.bold = True

    document.add_paragraph(
        exam.instructions
    )

    document.add_paragraph()

    # Questions
    for index, question in enumerate(
        exam.questions,
        start=1
    ):

        number = question.get(
            "number",
            index
        )

        question_text = question.get(
            "question",
            ""
        )

        paragraph = document.add_paragraph()

        run = paragraph.add_run(
            f"{number}. {question_text}"
        )

        run.font.size = Pt(11)

        choices = question.get(
            "choices",
            {}
        )

        for letter in ["A", "B", "C", "D"]:

            document.add_paragraph(
                f"   {letter}. {choices.get(letter, '')}"
            )

        document.add_paragraph()

    # Answer key page
    document.add_page_break()

    answer_title = document.add_paragraph()

    answer_run = answer_title.add_run(
        "ANSWER KEY"
    )

    answer_run.bold = True
    answer_run.font.size = Pt(16)

    for item in exam.answer_key:

        paragraph = document.add_paragraph()

        paragraph.add_run(
            f"{item['number']}. {item['answer']}"
        ).bold = True

        if item.get("explanation"):

            document.add_paragraph(
                f"Explanation: {item['explanation']}"
            )

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return StreamingResponse(
        output,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": (
                'attachment; filename="Mak-AI-Exam.docx"'
            )
        },
    )